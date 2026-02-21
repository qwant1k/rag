"""
Ingestion пайплайн — загрузка, парсинг и индексация документов в ChromaDB.
Поддерживаемые форматы: PDF (PyMuPDF), DOCX (python-docx), DOC (pywin32 COM), TXT.
"""

import logging
import hashlib
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
import pytesseract
import io
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

from backend.config import (
    DOCUMENTS_DIR,
    CHROMA_DB_DIR,
    EMBEDDING_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    CHROMA_COLLECTION_NAME,
    TESSERACT_CMD,
)

logger = logging.getLogger("rag-chatbot")

# Указываем путь к tesseract.exe для Windows
pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD


def normalize_text(text: str) -> str:
    """
    Базовая очистка текста перед индексацией:
    - исправление переносов слов через дефис,
    - удаление служебных unicode-символов,
    - нормализация пробелов и переносов.
    """
    if not text:
        return ""

    cleaned = text.replace("\r\n", "\n").replace("\r", "\n")
    cleaned = cleaned.replace("\u00ad", "")  
    cleaned = cleaned.replace("\xa0", " ")   
    cleaned = re.sub(r"([A-Za-zА-Яа-яЁё])-\n([A-Za-zА-Яа-яЁё])", r"\1\2", cleaned)
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def should_use_ocr(text: str) -> bool:
    """
    Для некоторых PDF текстовый слой формально есть, но он "пустой" для RAG
    (мало букв/мусор). В этом случае пробуем OCR.
    """
    if not text or not text.strip():
        return True

    letters = sum(1 for ch in text if ch.isalpha())
    return letters < 40


def iter_docx_blocks(doc: DocxDocument):
    """Итератор по блокам DOCX в порядке документа: абзацы и таблицы."""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def get_relative_source(file_path: Path) -> str:
    """
    Вычисляет относительный путь файла от папки DOCUMENTS_DIR.
    Например: documents/договоры/2024/договор1.pdf → договоры/2024/договор1.pdf
    Если файл вне DOCUMENTS_DIR — возвращает просто имя файла.
    """
    try:
        return str(file_path.relative_to(DOCUMENTS_DIR)).replace("\\", "/")
    except ValueError:
        return file_path.name


def get_embeddings() -> HuggingFaceEmbeddings:
    """Инициализация модели эмбеддингов (all-MiniLM-L6-v2)."""
    logger.info(f"Загрузка модели эмбеддингов: {EMBEDDING_MODEL}")
    return HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def get_vectorstore(embeddings: Optional[HuggingFaceEmbeddings] = None) -> Chroma:
    """Получение экземпляра ChromaDB vectorstore."""
    if embeddings is None:
        embeddings = get_embeddings()
    return Chroma(
        collection_name=CHROMA_COLLECTION_NAME,
        embedding_function=embeddings,
        persist_directory=str(CHROMA_DB_DIR),
    )


# === Парсеры документов ===

def ocr_page(page) -> str:
    """OCR одной страницы PDF через pytesseract (для сканированных документов)."""
    try:
        pix = page.get_pixmap(dpi=300)
        img_bytes = pix.tobytes("png")
        image = Image.open(io.BytesIO(img_bytes))
        text = pytesseract.image_to_string(image, lang="rus+eng+kaz")
        return text.strip()
    except Exception as e:
        logger.warning(f"OCR ошибка на странице: {e}")
        return ""


def parse_pdf(file_path: Path) -> list[Document]:
    """
    Парсинг PDF файла через PyMuPDF.
    Если страница не содержит текстового слоя — применяется OCR через pytesseract.
    """
    documents = []
    ocr_used = False
    try:
        pdf = fitz.open(str(file_path))
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            # Пробуем извлечь текстовый слой
            text = normalize_text(page.get_text("text"))
            # Если текста нет/мусорный текст — пробуем OCR
            if should_use_ocr(text):
                ocr_text = normalize_text(ocr_page(page))
                if ocr_text:
                    text = ocr_text
                    ocr_used = True
            if text:
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": get_relative_source(file_path),
                            "page": page_num + 1,
                            "upload_date": datetime.now().isoformat(),
                        },
                    )
                )
        pdf.close()
        method = "OCR" if ocr_used else "текстовый слой"
        logger.info(f"PDF '{get_relative_source(file_path)}': извлечено {len(documents)} страниц ({method})")
    except Exception as e:
        logger.error(f"Ошибка при парсинге PDF '{get_relative_source(file_path)}': {e}")
    return documents


def parse_docx(file_path: Path) -> list[Document]:
    """Парсинг DOCX файла через python-docx. Возвращает список Document."""
    documents = []
    try:
        doc = DocxDocument(str(file_path))
        blocks: list[str] = []

        for block in iter_docx_blocks(doc):
            if isinstance(block, Paragraph):
                text = normalize_text(block.text)
                if text:
                    blocks.append(text)
            elif isinstance(block, Table):
                for row in block.rows:
                    cells = [normalize_text(cell.text) for cell in row.cells]
                    cells = [cell for cell in cells if cell]
                    if cells:
                        blocks.append(" | ".join(cells))

        full_text = "\n".join(blocks)
        if full_text.strip():
            documents.append(
                Document(
                    page_content=full_text,
                    metadata={
                        "source": get_relative_source(file_path),
                        "page": 1,
                        "upload_date": datetime.now().isoformat(),
                    },
                )
            )
        logger.info(f"DOCX '{get_relative_source(file_path)}': извлечено {len(full_text)} символов")
    except Exception as e:
        logger.error(f"Ошибка при парсинге DOCX '{get_relative_source(file_path)}': {e}")
    return documents


def parse_txt(file_path: Path) -> list[Document]:
    """Парсинг TXT файла. Явно указываем кодировку UTF-8."""
    documents = []
    try:
        text = normalize_text(file_path.read_text(encoding="utf-8"))
        if text.strip():
            documents.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": get_relative_source(file_path),
                        "page": 1,
                        "upload_date": datetime.now().isoformat(),
                    },
                )
            )
        logger.info(f"TXT '{get_relative_source(file_path)}': извлечено {len(text)} символов")
    except UnicodeDecodeError:
        # Попытка прочитать с кодировкой cp1251 (частая на Windows)
        try:
            text = normalize_text(file_path.read_text(encoding="cp1251"))
            if text.strip():
                documents.append(
                    Document(
                        page_content=text,
                        metadata={
                            "source": get_relative_source(file_path),
                            "page": 1,
                            "upload_date": datetime.now().isoformat(),
                        },
                    )
                )
            logger.info(f"TXT '{get_relative_source(file_path)}': прочитан в cp1251, {len(text)} символов")
        except Exception as e:
            logger.error(f"Ошибка кодировки TXT '{get_relative_source(file_path)}': {e}")
    except Exception as e:
        logger.error(f"Ошибка при парсинге TXT '{get_relative_source(file_path)}': {e}")
    return documents


def parse_doc(file_path: Path) -> list[Document]:
    """
    Парсинг DOC файла (старый формат Word) через pywin32 COM.
    Требуется Windows с установленным MS Word.
    """
    documents = []
    try:
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(str(file_path.resolve()))
            text = doc.Content.Text
            doc.Close(False)
            word.Quit()
        finally:
            pythoncom.CoUninitialize()

        if text.strip():
            documents.append(
                Document(
                    page_content=normalize_text(text),
                    metadata={
                        "source": get_relative_source(file_path),
                        "page": 1,
                        "upload_date": datetime.now().isoformat(),
                    },
                )
            )
        logger.info(f"DOC '{get_relative_source(file_path)}': извлечено {len(text)} символов")
    except ImportError:
        logger.error("Для чтения .doc файлов требуется pywin32: pip install pywin32")
    except Exception as e:
        logger.error(f"Ошибка при парсинге DOC '{get_relative_source(file_path)}': {e}")
    return documents


def parse_file(file_path: Path) -> list[Document]:
    """Универсальный парсер — определяет тип файла и вызывает нужный парсер."""
    suffix = file_path.suffix.lower()
    parsers = {
        ".pdf": parse_pdf,
        ".docx": parse_docx,
        ".doc": parse_doc,
        ".txt": parse_txt,
    }
    parser = parsers.get(suffix)
    if parser is None:
        logger.warning(f"Неподдерживаемый формат файла: {file_path.name}")
        return []
    return parser(file_path)


def split_documents(documents: list[Document]) -> list[Document]:
    """Разбивка документов на чанки заданного размера с перекрытием."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ": ", " ", ""],
    )
    chunks = text_splitter.split_documents(documents)
    logger.info(f"Разбивка: {len(documents)} документов → {len(chunks)} чанков")
    return chunks


def generate_chunk_id(chunk: Document, index: int) -> str:
    """Генерация уникального ID для чанка на основе содержимого и метаданных."""
    content = f"{chunk.metadata.get('source', '')}_{chunk.metadata.get('page', '')}_{index}_{chunk.page_content[:100]}"
    return hashlib.md5(content.encode("utf-8")).hexdigest()


def delete_document_from_db(filename: str, vectorstore: Optional[Chroma] = None) -> int:
    """
    Удаление всех чанков документа из ChromaDB по имени файла.
    Возвращает количество удалённых чанков.
    """
    if vectorstore is None:
        vectorstore = get_vectorstore()

    # Получаем коллекцию напрямую для фильтрации
    collection = vectorstore._collection
    results = collection.get(where={"source": filename})

    if results and results["ids"]:
        count = len(results["ids"])
        collection.delete(ids=results["ids"])
        logger.info(f"Удалено {count} чанков документа '{filename}' из ChromaDB")
        return count

    logger.info(f"Документ '{filename}' не найден в ChromaDB")
    return 0


def ingest_file(file_path: Path, vectorstore: Optional[Chroma] = None) -> int:
    """
    Полный пайплайн загрузки одного файла:
    1. Парсинг файла
    2. Разбивка на чанки
    3. Удаление старых чанков этого файла (дедупликация)
    4. Добавление новых чанков в ChromaDB
    Возвращает количество добавленных чанков.
    """
    if vectorstore is None:
        vectorstore = get_vectorstore()

    source_name = get_relative_source(file_path)
    logger.info(f"Начало обработки файла: {source_name}")

    # Шаг 1: Парсинг
    documents = parse_file(file_path)
    if not documents:
        logger.warning(f"Файл '{source_name}' не содержит текста или не распознан")
        return 0

    # Шаг 2: Разбивка на чанки
    chunks = split_documents(documents)
    if not chunks:
        logger.warning(f"Файл '{source_name}': после разбивки нет чанков")
        return 0

    # Шаг 3: Дедупликация — удаляем старые чанки этого файла (по относительному пути)
    delete_document_from_db(source_name, vectorstore)

    # Шаг 4: Генерация ID и добавление в ChromaDB
    ids = [generate_chunk_id(chunk, i) for i, chunk in enumerate(chunks)]
    vectorstore.add_documents(documents=chunks, ids=ids)

    logger.info(f"✅ Файл '{source_name}': добавлено {len(chunks)} чанков в ChromaDB")
    return len(chunks)


def ingest_directory(directory: Optional[Path] = None) -> dict:
    """
    Рекурсивная обработка всех файлов из указанной директории и вложенных папок.
    Использует Path.rglob() для обхода на любую глубину вложенности.
    Возвращает словарь {относительный_путь: количество_чанков}.
    """
    if directory is None:
        directory = DOCUMENTS_DIR

    if not directory.exists():
        logger.error(f"Директория не существует: {directory}")
        return {}

    # Поддерживаемые расширения — рекурсивный поиск через rglob
    extensions = {".pdf", ".docx", ".doc", ".txt"}
    files = [
        f for f in directory.rglob("*")
        if f.is_file()
        and f.suffix.lower() in extensions
        and not f.name.startswith("~$")  # Игнорируем временные файлы Word
    ]

    if not files:
        logger.warning(f"В директории '{directory}' (включая вложенные) нет поддерживаемых файлов")
        return {}

    logger.info(f"Найдено {len(files)} файлов для обработки (рекурсивный обход)")

    # Создаём vectorstore и embeddings один раз для всех файлов
    embeddings = get_embeddings()
    vectorstore = get_vectorstore(embeddings)

    results = {}
    for file_path in sorted(files):
        source_name = get_relative_source(file_path)
        try:
            count = ingest_file(file_path, vectorstore)
            results[source_name] = count
        except Exception as e:
            logger.error(f"Ошибка при обработке '{source_name}': {e}")
            results[source_name] = 0

    total_chunks = sum(results.values())
    logger.info(f"🎉 Обработка завершена: {len(results)} файлов, {total_chunks} чанков всего")
    return results


def get_indexed_documents() -> list[dict]:
    """
    Получение списка всех проиндексированных документов из ChromaDB.
    Возвращает список словарей с информацией о каждом документе.
    """
    try:
        vectorstore = get_vectorstore()
        collection = vectorstore._collection
        all_data = collection.get(include=["metadatas"])

        if not all_data or not all_data["metadatas"]:
            return []

        # Группируем по имени файла
        doc_info = {}
        for metadata in all_data["metadatas"]:
            source = metadata.get("source", "unknown")
            if source not in doc_info:
                doc_info[source] = {
                    "filename": source,
                    "chunks_count": 0,
                    "pages": set(),
                    "upload_date": metadata.get("upload_date", ""),
                }
            doc_info[source]["chunks_count"] += 1
            page = metadata.get("page")
            if page:
                doc_info[source]["pages"].add(page)

        # Преобразуем set в sorted list для JSON сериализации
        result = []
        for info in doc_info.values():
            info["pages"] = sorted(info["pages"])
            result.append(info)

        return result
    except Exception as e:
        logger.error(f"Ошибка при получении списка документов: {e}")
        return []


# === Точка входа для запуска из командной строки ===
if __name__ == "__main__":
    logger.info("=" * 60)
    logger.info("Запуск индексации документов из папки /documents")
    logger.info("=" * 60)

    results = ingest_directory()

    if results:
        print("\n📊 Результаты индексации:")
        print("-" * 40)
        for filename, count in results.items():
            status = "✅" if count > 0 else "❌"
            print(f"  {status} {filename}: {count} чанков")
        print("-" * 40)
        print(f"  Всего чанков: {sum(results.values())}")
    else:
        print("\n⚠️ Нет файлов для индексации. Положите PDF, DOCX или TXT файлы в папку /documents")
